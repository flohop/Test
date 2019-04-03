from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import wikipedia
import os
import sys
import shutil

import word_to_pdf
from wiki_api import WikiPage
import email_sender


class WordWriter():
    """Take all the infos from wiki_api and write them in a word-file"""
    def __init__(self, page_name, language="de"):
        try:
            self.wiki_page = WikiPage(page_name, language)
        except wikipedia.exceptions.DisambiguationError:
            print("ERROR: Please specify your search")
            sys.exit()

        self.pic = self.wiki_page.download_picture()
        self.document = Document()

    def add_title(self):
        """Add the heading to Word-file"""
        title = self.document.add_heading("", level=1)
        run = title.add_run()
        font = run.font
        font.color.rgb = RGBColor(0, 0, 0)
        run.text = str(self.wiki_page.get_title())
        run.underline = True
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def add_logo(self):
        """Add the image to the Word-file"""
        """error somewhere(napoleon)"""
        try:
            self.document.add_picture(str(self.wiki_page.download_picture()), width=Inches(1.25))
            last_paragraph = self.document.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        except IndexError:
            print("No image found")
            pass

    def add_summary(self):
        """Add the summary to the Word-file"""
        summary = self.document.add_paragraph(str(self.wiki_page.get_summary()), style="Footer")
        summary.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_s = summary.add_run()
        run_s.add_break()
        run_s.add_break()
        style = self.document.styles["Footer"]
        font = style.font
        font.size = Pt(14)

    def edit_content(self, fat_content, *args):
        """Cut out certain parts of the article"""
        cut_out_sections = []
        fat_content = fat_content.replace("===", "==")

        #  cut out the parts in *args
        for arg in args:
            cut_out_sections.append(str(arg))

        for item in cut_out_sections:

            cut_out_content = self.wiki_page.get_section(item)  # get the cut out content

            try:
                section_name = "== " + str(item) + " =="
                slim_content = fat_content.replace(cut_out_content, "")  # replace die cut out contents with " "
                slim_sections = slim_content.replace(section_name, "")
                fat_content = str(slim_sections).rstrip("\n")
            except TypeError:
                pass

        return fat_content.rstrip("\r")

    def add_content(self, content):
        """Add the content returned from edit_content to the Word-file"""
        content = self.document.add_paragraph(content, style='Quote')
        content.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        style = self.document.styles["Quote"]
        font = style.font
        font.size = Pt(14)


    def delete_pic_in_outer_folder(self):
        """Delete the pictures in the outer folder"""
        for file in os.listdir():
            if file.endswith(".jpg"):
                os.unlink(file)

    def get_doc_name(self):
        """Return name of the Wikipedia article"""
        return str(self.wiki_page.get_title())


    def move_word_files(self):
        """move the word files in the word_folder"""
        destination_file = os.getcwd() + "\word_files"
        for file in os.listdir():
            if file.endswith(".docx"):
                try:
                    shutil.move(file, destination_file)
                except shutil.Error:
                    pass
        for file in os.listdir():
            if file.endswith('.docx'):
                os.unlink(file)

    def make_send_pdf(self):
        """Create the pdf and send it via email"""
        slim_content = edit_content(page.get_content(page.get_summary()), "Weblinks", "Literatur", "Einzelnachweise",
                                    "Sonstiges", "Filme", "Auszeichnungen", "Filmdokumentationen", "Anmerkungen",
                                    "Biografien",
                                    "Weitere Texte")  # cuts out these sections from the article

        to_mail = 'send_to_this_mail@gmail.com'  # testing email
        try:
            print("Progress:")
            # create the word page
            add_title()
            print("Added title...")
            add_logo()
            print("Added image...")
            add_summary()
            print("Added summary...")
            add_content(slim_content)
            print("Added content...")
            document.save(get_doc_name() + ".docx")
            print("Saving file...")
            delete_pic_in_outer_folder()
            word_to_pdf.docx_to_pdf(get_doc_name() + ".docx", get_doc_name())
            move_word_files()
            print("Sending email...")
            email_sender.send_email(to_mail, get_doc_name() + ".pdf", str(page.get_title()))  # uncomment to send email
            word_to_pdf.move_pdf_to_folder()
            print("===Finished===")
        except PermissionError:
            print(
                "Bitte schließen sie Microsoft Word und versuchen sie es erneut.")  # "Please close Microsoft Word and try again"

